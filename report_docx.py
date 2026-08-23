from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PURPLE = RGBColor(109, 67, 219)
DARK = RGBColor(43, 45, 58)
MUTED = RGBColor(111, 115, 132)
LIGHT = "F2EEFF"


def _font(run, size: float, color: RGBColor = DARK, bold: bool = False) -> None:
    run.font.name = "Malgun Gothic"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run("Page ")
    _font(run, 8.5, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def _set_cell_margins(cell, top=120, start=160, bottom=120, end=160) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def build_report_docx(report: str, keywords: Iterable[str], item_count: int,
                      lookback_hours: int, generated_at: datetime | None = None) -> bytes:
    generated_at = generated_at or datetime.now()
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.text = "GOOGLE NEWS  |  KEYWORD INTELLIGENCE"
    _font(header.runs[0], 8.5, MUTED, True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_page_number(footer)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    _font(kicker.add_run("MARKET MONITORING BRIEF"), 9, PURPLE, True)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    _font(title.add_run("Google 뉴스 키워드 리포트"), 23, DARK, True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _font(subtitle.add_run("검색 신호를 실행 가능한 이슈와 근거로 정리한 업무용 보고서"), 11, MUTED)

    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    values = [("생성 시각", generated_at.strftime("%Y.%m.%d %H:%M")),
              ("분석 기사", f"{item_count}건"), ("수집 범위", f"최근 {lookback_hours}시간"),
              ("키워드", ", ".join(keywords) or "미지정")]
    for index, (label, value) in enumerate(values):
        cell = table.cell(0, index)
        _shade(cell, LIGHT)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _font(p.add_run(label + "\n"), 8, PURPLE, True)
        _font(p.add_run(value), 9.2, DARK, True)
    document.add_paragraph().paragraph_format.space_after = Pt(0)

    for raw_line in report.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        marker = line[0] if line[0] in "□-·" else "·"
        text = line[1:].strip() if line[0] in "□-·" else line
        p = document.add_paragraph()
        p.paragraph_format.keep_together = True
        if marker == "□":
            p.paragraph_format.space_before = Pt(13)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.keep_with_next = True
            _font(p.add_run("□  "), 14, PURPLE, True)
            _font(p.add_run(text), 14, DARK, True)
        elif marker == "-":
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.12
            _font(p.add_run("-  "), 10.5, PURPLE, True)
            _font(p.add_run(text), 10.5, DARK)
        else:
            p.paragraph_format.left_indent = Inches(0.46)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.1
            _font(p.add_run("·  "), 9.5, PURPLE, True)
            _font(p.add_run(text), 9.5, MUTED)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
